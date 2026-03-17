from __future__ import annotations

import io
import json
import os
import secrets
from datetime import datetime
from typing import Any, Optional

import pdfplumber
from docx import Document
from flask import Blueprint, jsonify, request, send_file
from docx import Document as DocxDocument

from ai_client import AIClient
from ai_helper import AIHelper, _safe_json_dumps
from career_graph import build_echarts_data, list_all_families
from context_memory import (
    add_key_fact,
    get as get_context,
    set_personality_diagnosis,
    set_profile_summary,
    set_regret_result,
    to_context_string,
)
from game_theory import regret_matching_from_profile
from knowledge import search_knowledge
from database import db
from models import Student, Job, MatchResult, Report, ChatMessage


api = Blueprint("api", __name__)

TOKENS: dict[str, int] = {}


def _get_bearer_token() -> str:
    auth = request.headers.get("Authorization", "")
    if auth.lower().startswith("bearer "):
        return auth.split(" ", 1)[1].strip()
    return ""


def _require_student() -> Student:
    token = _get_bearer_token()
    student_id = TOKENS.get(token)
    if not student_id:
        raise PermissionError("未登录或登录已过期")
    student = db.session.get(Student, student_id)
    if not student:
        raise PermissionError("用户不存在")
    return student


def _json_error(message: str, status: int = 400):
    return jsonify({"ok": False, "error": message}), status


def _read_docx_bytes(data: bytes) -> str:
    f = io.BytesIO(data)
    doc = Document(f)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _read_pdf_bytes(data: bytes) -> str:
    f = io.BytesIO(data)
    parts = []
    with pdfplumber.open(f) as pdf:
        for page in pdf.pages[:10]:
            txt = page.extract_text() or ""
            if txt.strip():
                parts.append(txt)
    return "\n".join(parts)


def _seed_jobs_if_empty(ai: AIHelper) -> None:
    if db.session.query(Job).count() > 0:
        return

    samples = [
        {
            "title": "数据分析实习生",
            "industry": "互联网/数据",
            "salary": "150-200/天",
            "requirements_text": "熟悉Excel/SQL，了解Python数据分析，良好沟通，能做可视化报表。",
        },
        {
            "title": "后端开发工程师（校招）",
            "industry": "互联网/软件",
            "salary": "12-20K",
            "requirements_text": "掌握Python/Java至少一种，熟悉Flask/Django或Spring，理解数据库与缓存，了解Linux与Git。",
        },
        {
            "title": "前端开发工程师（校招）",
            "industry": "互联网/软件",
            "salary": "12-20K",
            "requirements_text": "熟悉JavaScript/TypeScript，掌握React或Vue，了解工程化与性能优化，有作品集优先。",
        },
        {
            "title": "产品助理",
            "industry": "互联网/产品",
            "salary": "10-16K",
            "requirements_text": "具备用户调研与需求分析能力，良好沟通与推动，能写PRD，会用原型工具。",
        },
        {
            "title": "算法工程师（NLP方向）",
            "industry": "AI/算法",
            "salary": "18-35K",
            "requirements_text": "熟悉机器学习/深度学习，掌握PyTorch，了解NLP/LLM，具备论文复现与工程落地能力。",
        },
    ]

    for s in samples:
        job = Job(
            title=s["title"],
            industry=s.get("industry"),
            salary=s.get("salary"),
            requirements_text=s.get("requirements_text"),
            job_profile_json=_safe_json_dumps(ai.generate_job_profile(s)),
        )
        db.session.add(job)

    db.session.commit()


def _get_ai() -> AIHelper:
    return AIHelper(client=AIClient())


@api.errorhandler(PermissionError)
def _handle_perm(e):
    return _json_error(str(e), status=401)


@api.route("/login", methods=["POST"])
def login():
    payload = request.get_json(silent=True) or {}
    username = (payload.get("username") or "").strip()
    password = (payload.get("password") or "").strip()

    if username != "admin" or password != "123456":
        return _json_error("用户名或密码错误（测试账号：admin / 123456）", status=401)

    student = db.session.query(Student).filter_by(username=username).one_or_none()
    if not student:
        student = Student(username=username, name="测试用户", major="计算机相关")
        db.session.add(student)
        db.session.commit()

    token = secrets.token_urlsafe(24)
    TOKENS[token] = student.id

    return jsonify({"ok": True, "token": token, "student_id": student.id, "student": student.to_dict()})


@api.route("/upload_resume", methods=["POST"])
def upload_resume():
    student = _require_student()
    ai = _get_ai()
    _seed_jobs_if_empty(ai)

    manual = {
        "major": (request.form.get("major") or "").strip(),
        "skills": (request.form.get("skills") or "").strip(),
        "certs": (request.form.get("certs") or "").strip(),
        "internships": (request.form.get("internships") or "").strip(),
    }
    resume_text = (request.form.get("resume_text") or "").strip()

    file = request.files.get("resume_file")
    if file and file.filename:
        filename = file.filename.lower()
        data = file.read()
        if filename.endswith(".docx"):
            resume_text = (resume_text + "\n" + _read_docx_bytes(data)).strip()
        elif filename.endswith(".pdf"):
            resume_text = (resume_text + "\n" + _read_pdf_bytes(data)).strip()
        elif filename.endswith(".jpg") or filename.endswith(".jpeg") or filename.endswith(".png"):
            # OCR 依赖较重，这里先存档提示；企业需要 OCR 时可接入 PaddleOCR/Tesseract
            resume_text = (resume_text + "\n" + "（已上传图片简历：当前版本未启用OCR解析，请在手动录入区补充关键信息。）").strip()
        else:
            return _json_error("仅支持 docx/pdf/jpg/png 文件")

    if not resume_text:
        return _json_error("请上传简历文件或手动录入简历信息")

    parsed = ai.parse_resume(resume_text, manual=manual)

    if manual.get("major"):
        student.major = manual["major"]
    student.resume_raw_text = resume_text
    student.resume_parsed_json = _safe_json_dumps(parsed)
    db.session.commit()

    # 轻量级记忆：画像摘要，供 Agent 上下文联系
    summary = f"专业:{student.major or '未知'}；技能:{','.join((parsed.get('skills') or [])[:8])}；竞争力:{parsed.get('competitiveness_score','—')}"
    set_profile_summary(student.id, summary)

    # 写入一条 AI 引导消息（回到 chat.html 不丢失）
    db.session.add(
        ChatMessage(
            student_id=student.id,
            role="ai",
            content="简历解析完成！你的个人能力画像已生成，点击【性格测评】完善你的职业偏好吧～",
        )
    )
    db.session.commit()

    return jsonify({"ok": True, "student_id": student.id, "resume_parsed": parsed})


@api.route("/submit_test", methods=["POST"])
def submit_test():
    student = _require_student()
    payload = request.get_json(silent=True) or {}
    student.personality_test_json = _safe_json_dumps(payload)
    db.session.commit()

    # 性格分析 → 短板诊断 + 岗位适配，写入轻量级记忆（上下文联系）
    try:
        resume_profile = json.loads(student.resume_parsed_json or "{}")
    except Exception:
        resume_profile = {}
    personality_test = payload
    ai = _get_ai()
    diag = ai.analyze_personality_for_jobs(resume_profile, personality_test)
    set_personality_diagnosis(student.id, diag)

    db.session.add(
        ChatMessage(
            student_id=student.id,
            role="ai",
            content="职业偏好已完善，我会结合你的简历和测评做精准人岗匹配，结果会在【岗位列表】中标注匹配度哦～",
        )
    )
    db.session.commit()
    return jsonify({"ok": True, "student_id": student.id, "personality_diagnosis": diag})


@api.route("/jobs", methods=["GET"])
def get_jobs():
    student: Optional[Student] = None
    try:
        student = _require_student()
    except Exception:
        student = None

    ai = _get_ai()
    _seed_jobs_if_empty(ai)

    q = (request.args.get("q") or "").strip().lower()
    include_match = (request.args.get("include_match") or "").strip() in {"1", "true", "yes"}

    query = db.session.query(Job)
    if q:
        like = f"%{q}%"
        query = query.filter((Job.title.like(like)) | (Job.industry.like(like)) | (Job.requirements_text.like(like)))

    jobs = query.order_by(Job.id.asc()).limit(200).all()
    out = []

    student_profile = None
    if include_match and student and student.resume_parsed_json:
        try:
            student_profile = json.loads(student.resume_parsed_json)
        except Exception:
            student_profile = None

    for j in jobs:
        item = j.to_dict()
        if include_match and student_profile:
            try:
                job_profile = json.loads(j.job_profile_json or "{}")
            except Exception:
                job_profile = ai.generate_job_profile(item)
            m = ai.match(student_profile, job_profile)
            item["match_preview"] = {"score": m.get("score", 0), "reasoning": m.get("reasoning", "")}
        out.append(item)

    return jsonify({"ok": True, "jobs": out})


@api.route("/jobs/<int:job_id>", methods=["GET"])
def job_detail(job_id: int):
    job = db.session.get(Job, job_id)
    if not job:
        return _json_error("岗位不存在", status=404)
    return jsonify({"ok": True, "job": job.to_dict()})


@api.route("/me", methods=["GET"])
def me():
    student = _require_student()
    data = student.to_dict()
    try:
        data["resume_parsed"] = json.loads(student.resume_parsed_json or "{}")
    except Exception:
        data["resume_parsed"] = {}
    try:
        data["personality_test"] = json.loads(student.personality_test_json or "{}")
    except Exception:
        data["personality_test"] = {}
    return jsonify({"ok": True, "student": data})


@api.route("/match_jobs", methods=["POST"])
def match_jobs():
    student = _require_student()
    ai = _get_ai()
    _seed_jobs_if_empty(ai)

    payload = request.get_json(silent=True) or {}
    goal = (payload.get("goal") or "").strip()

    if not student.resume_parsed_json:
        return _json_error("请先上传/录入简历并解析")

    try:
        student_profile = json.loads(student.resume_parsed_json)
    except Exception:
        return _json_error("简历解析数据损坏，请重新上传解析")

    jobs = db.session.query(Job).order_by(Job.id.asc()).all()
    results = []

    # 清理旧匹配（演示版：每次重算覆盖）
    db.session.query(MatchResult).filter(MatchResult.student_id == student.id).delete()
    db.session.commit()

    for j in jobs:
        job_dict = j.to_dict()
        try:
            job_profile = json.loads(j.job_profile_json or "{}")
        except Exception:
            job_profile = ai.generate_job_profile(job_dict)
            j.job_profile_json = _safe_json_dumps(job_profile)
            db.session.add(j)

        matched = ai.match(student_profile, job_profile)
        dim_full = matched.get("dimension_scores") or {}
        dim_4 = matched.get("dimension_scores_4") or {}
        mr = MatchResult(
            student_id=student.id,
            job_id=j.id,
            score=float(matched.get("score") or 0),
            dimension_scores_json=_safe_json_dumps({"dimension_scores": dim_full, "dimension_scores_4": dim_4}),
            gap_analysis_json=_safe_json_dumps(matched.get("gap_analysis") or {}),
            reasoning=matched.get("reasoning") or "",
        )
        db.session.add(mr)
        results.append(
            {
                "job_id": j.id,
                "job_title": j.title,
                "industry": j.industry,
                "salary": j.salary,
                "score": mr.score,
                "dimension_scores": dim_full,
                "dimension_scores_4": dim_4,
                "gap_analysis": matched.get("gap_analysis") or {},
                "reasoning": mr.reasoning,
            }
        )

    db.session.commit()

    results.sort(key=lambda x: x["score"], reverse=True)
    return jsonify({"ok": True, "student_id": student.id, "goal": goal, "matches": results})


@api.route("/generate_report", methods=["POST"])
def generate_report():
    student = _require_student()
    ai = _get_ai()
    _seed_jobs_if_empty(ai)

    payload = request.get_json(silent=True) or {}
    goal = (payload.get("goal") or "").strip()

    if not student.resume_parsed_json:
        return _json_error("请先上传/录入简历并解析")

    # 取已有匹配或先计算
    match_rows = (
        db.session.query(MatchResult)
        .filter(MatchResult.student_id == student.id)
        .order_by(MatchResult.score.desc())
        .limit(5)
        .all()
    )
    if not match_rows:
        # 自动计算一次
        _ = match_jobs()
        match_rows = (
            db.session.query(MatchResult)
            .filter(MatchResult.student_id == student.id)
            .order_by(MatchResult.score.desc())
            .limit(5)
            .all()
        )

    top_matches = []
    for mr in match_rows:
        job = db.session.get(Job, mr.job_id)
        try:
            dims = json.loads(mr.dimension_scores_json or "{}")
            dim_scores = dims.get("dimension_scores") or dims
            dim_scores_4 = dims.get("dimension_scores_4") or {}
        except Exception:
            dim_scores = {}
            dim_scores_4 = {}
        top_matches.append(
            {
                "job_id": mr.job_id,
                "job_title": job.title if job else f"job#{mr.job_id}",
                "score": mr.score,
                "gap_analysis": json.loads(mr.gap_analysis_json or "{}"),
                "dimension_scores": dim_scores,
                "dimension_scores_4": dim_scores_4,
                "reasoning": mr.reasoning,
            }
        )

    student_dict = student.to_dict()
    content = ai.generate_report(student_dict, top_matches, goal=goal)

    from ai_helper import report_completeness_check
    check = report_completeness_check(content)

    report = Report(student_id=student.id, content=content)
    db.session.add(report)
    db.session.commit()

    return jsonify(
        {
            "ok": True,
            "report_id": report.id,
            "content": content,
            "download_url": f"/download_report/{report.id}",
            "download_docx_url": f"/download_report_docx/{report.id}",
            "completeness_check": check,
            "created_at": datetime.utcnow().isoformat(),
        }
    )


@api.route("/career_graph", methods=["GET"])
def career_graph():
    """岗位关联图谱：垂直晋升 + 换岗路径。至少 5 个岗位族，每族不少于 2 条换岗路径。"""
    title = (request.args.get("title") or "").strip()
    if title:
        data = build_echarts_data(title)
        return jsonify({"ok": True, "graph": data, "families": list_all_families()})
    return jsonify({"ok": True, "families": list_all_families(), "graph": build_echarts_data(None)})


@api.route("/chat/history", methods=["GET"])
def chat_history():
    student = _require_student()
    limit = int(request.args.get("limit") or 50)
    rows = (
        db.session.query(ChatMessage)
        .filter(ChatMessage.student_id == student.id)
        .order_by(ChatMessage.id.asc())
        .limit(max(1, min(limit, 300)))
        .all()
    )
    return jsonify({"ok": True, "messages": [r.to_dict() for r in rows]})


def _offline_agent_reply(student: Student, user_text: str) -> str:
    # 离线兜底：给出引导+可执行建议（不联网）
    if not student.resume_parsed_json:
        return "我可以先帮你解析简历生成能力画像。请点击顶部【简历上传】上传/录入简历信息。"
    if not student.personality_test_json:
        return "我已拿到你的简历画像。建议再做一次【性格测评】补充职业偏好，我会基于两者做更精准的人岗匹配。"
    t = (user_text or "").strip()
    if "报告" in t:
        return "你可以点击顶部【报告生成】一键生成职业生涯发展报告（含目标与行动计划），并支持下载。"
    if "岗位" in t or "匹配" in t:
        return "我建议你先在【岗位列表】查看匹配度较高的岗位，并进入【匹配结果】页查看 TOP5、雷达图与差距清单。"
    return "你可以问我：市场岗位能力要求、我的优势与短板、推荐岗位与差距、以及分阶段行动计划。"


@api.route("/chat/send", methods=["POST"])
def chat_send():
    student = _require_student()
    payload = request.get_json(silent=True) or {}
    user_text = (payload.get("text") or "").strip()
    if not user_text:
        return _json_error("消息不能为空")

    # 存用户消息
    db.session.add(ChatMessage(student_id=student.id, role="user", content=user_text))
    db.session.commit()

    ai = _get_ai()
    client = ai.client

    # 轻量级记忆：先用 LLM 做一句话关键信息摘要，失败则用关键词匹配兜底
    try:
        summary = client.chat_text(
            "你是信息抽取助手，请用不超过一句话的中文提炼这段话的关键信息。",
            f"用户原话：{user_text}",
        )
        if summary:
            add_key_fact(student.id, summary.strip())
    except Exception:
        key_phrases = ["成绩", "内向", "外向", "压力", "家庭", "考研", "保研", "大厂", "国企", "实习", "项目", "证书", "迷茫", "焦虑"]
        for p in key_phrases:
            if p in user_text:
                add_key_fact(student.id, f"用户提及：{user_text[:80]}...")
                break

    # 博弈模型：Regret Matching 计算最优路径（大厂/国企/保研）
    combined = {"dimensions": {}, "personality": {}}
    try:
        rp = json.loads(student.resume_parsed_json or "{}")
        combined["dimensions"] = rp.get("dimensions") or {}
        combined["competitiveness_score"] = rp.get("competitiveness_score") or 50
    except Exception:
        pass
    try:
        pt = json.loads(student.personality_test_json or "{}")
        s = pt.get("scores") or {}
        combined["personality"] = {k: v for k, v in s.items()}
        combined["career_interest"] = pt.get("career_interest") or {}
        combined["resources"] = pt.get("resources") or []
        combined["constraints"] = pt.get("constraints") or {}
    except Exception:
        pass

    # 即时偏好：更想稳定 / 不想出国 等（影响博弈参数）
    prefs: dict[str, bool] = {}
    if "稳定" in user_text or "体制" in user_text or "编制" in user_text:
        prefs["prefer_stable"] = True
    if "不想出国" in user_text or "不想海外" in user_text or "留在国内" in user_text:
        prefs["avoid_abroad"] = True
    combined["user_prefs"] = prefs
    try:
        regret_res = regret_matching_from_profile(combined)
        set_regret_result(student.id, regret_res)
    except Exception:
        regret_res = None

    # 拼接上下文：记忆模块 + 博弈结果 + 学生画像 + 最近对话
    history_rows = (
        db.session.query(ChatMessage)
        .filter(ChatMessage.student_id == student.id)
        .order_by(ChatMessage.id.desc())
        .limit(20)
        .all()
    )
    history = list(reversed(history_rows))

    # 简易联网搜索 + 本地知识库：当用户询问“市场要求/竞赛/趋势/怎么学”等，补充检索摘要
    search_block = ""
    keywords = ["市场", "趋势", "要求", "竞赛", "证书", "学习路线", "怎么学", "什么是", "是什么"]
    if any(k in user_text for k in keywords):
        results = client.web_search(user_text, limit=5) if hasattr(client, "web_search") else []
        if results:
            lines = [f"- {r.get('title','')}（{r.get('url','')}）" for r in results]
            search_block = "【联网检索参考】\n" + "\n".join(lines) + "\n"
        # 本地知识库（轻量 RAG）
        local_snips = search_knowledge(user_text, max_chunks=3)
        if local_snips:
            ls = [s.get("snippet", "") for s in local_snips]
            search_block += "【本地知识库参考】\n" + "\n---\n".join(ls[:3]) + "\n"

    system_prompt = (
        "你是温暖、专业的大学生 AI 职业规划助手，像一位关心你的学长/学姐。"
        "你的目标：1) 帮学生快速了解就业市场/岗位能力要求；2) 准确分析就业能力与意愿；"
        "3) 给出明确可操作的建议。要求：语气亲切、有共情，同时专业、可执行；"
        "若系统提供了【性格与短板诊断】或【博弈路径建议】，务必结合这些信息综合给出最优建议。"
        "输出中文。\n"
    )

    memory_block = to_context_string(student.id)
    regret_block = ""
    if regret_res:
        regret_block = f"\n【博弈路径建议（Regret Matching）】最优路径：{regret_res.best_action}；{regret_res.recommendation}"

    context = {
        "student": student.to_dict(),
        "resume_parsed_json": student.resume_parsed_json,
        "personality_test_json": student.personality_test_json,
        "search": search_block,
        "recent_messages": [{"role": m.role, "content": m.content} for m in history],
    }
    user_prompt = (
        f"{search_block}"
        f"{memory_block}\n{regret_block}\n\n"
        f"【学生信息与画像】\n{_safe_json_dumps(context)}\n\n"
        f"【本轮用户问题】\n{user_text}\n\n"
        "请结合记忆、性格诊断、博弈路径建议与对话历史，亲切、专业地回答，并给出下一步操作引导。"
    )

    reply = ""
    try:
        reply = client.chat_text(system_prompt, user_prompt)
    except Exception:
        reply = _offline_agent_reply(student, user_text)

    db.session.add(ChatMessage(student_id=student.id, role="ai", content=reply))
    db.session.commit()

    return jsonify({"ok": True, "reply": reply})


@api.route("/download_report/<int:report_id>", methods=["GET"])
def download_report(report_id: int):
    student = _require_student()
    report = db.session.get(Report, report_id)
    if not report or report.student_id != student.id:
        return _json_error("报告不存在", status=404)

    data = report.content.encode("utf-8")
    return send_file(
        io.BytesIO(data),
        as_attachment=True,
        download_name=f"career_report_{report_id}.txt",
        mimetype="text/plain; charset=utf-8",
    )


@api.route("/download_report_docx/<int:report_id>", methods=["GET"])
def download_report_docx(report_id: int):
    student = _require_student()
    report = db.session.get(Report, report_id)
    if not report or report.student_id != student.id:
        return _json_error("报告不存在", status=404)

    doc = DocxDocument()
    for line in (report.content or "").splitlines():
        doc.add_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return send_file(
        bio,
        as_attachment=True,
        download_name=f"career_report_{report_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


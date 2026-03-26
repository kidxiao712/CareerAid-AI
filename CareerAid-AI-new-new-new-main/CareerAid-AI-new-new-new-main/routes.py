from __future__ import annotations

import os
os.environ['PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK'] = 'True'

import io
import json
import secrets
import hashlib
from datetime import datetime
from typing import Any, Optional
# from paddleocr import PaddleOCR
from PIL import Image

import pdfplumber
from docx import Document
from flask import Blueprint, jsonify, request, send_file
from docx import Document as DocxDocument

from ai_client import AIClient
from ai_helper import AIHelper, _safe_json_dumps
from career_graph import build_echarts_data, list_all_families, build_dynamic_echarts_data
from context_memory import (
    add_key_fact,
    get as get_context,
    set_personality_diagnosis,
    set_profile_summary,
    set_regret_result,
    to_context_string,
)
from game_theory import regret_matching_from_profile
from knowledge import search_knowledge
from database import db
from models import Student, Job, MatchResult, Report, ChatMessage, Feedback

import json
import time
from datetime import datetime
from cache import (
    cache_learning_plan,
    get_cached_learning_plan,
    cache_match_result,
    get_cached_match_result,
    cache_job_profile,
    get_cached_job_profile
)
from trend_analysis import analyze_industry_trends, analyze_job_relations


api = Blueprint("api", __name__)

TOKENS: dict[str, int] = {}


def _get_bearer_token() -> str:
    auth = request.headers.get("Authorization", "")
    if auth.lower().startswith("bearer "):
        return auth.split(" ", 1)[1].strip()
    return ""


def _require_student() -> Student:
    token = _get_bearer_token()
    student_id = TOKENS.get(token)
    if not student_id:
        raise PermissionError("未登录或登录已过期")
    student = db.session.get(Student, student_id)
    if not student:
        raise PermissionError("用户不存在")
    return student


def _require_admin() -> Student:
    """要求管理员权限"""
    student = _require_student()
    if student.username != "admin":
        raise PermissionError("需要管理员权限")
    return student


def _hash_password(password: str) -> str:
    """对密码进行哈希处理"""
    return hashlib.sha256(password.encode()).hexdigest()


def _json_error(message: str, status: int = 400):
    return jsonify({"ok": False, "error": message}), status


def _read_docx_bytes(data: bytes) -> str:
    f = io.BytesIO(data)
    doc = Document(f)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _read_pdf_bytes(data: bytes) -> str:
    f = io.BytesIO(data)
    parts = []
    with pdfplumber.open(f) as pdf:
        for page in pdf.pages[:10]:
            txt = page.extract_text() or ""
            if txt.strip():
                parts.append(txt)
    return "\n".join(parts)


def _seed_jobs_if_empty(ai: AIHelper) -> None:
    """检查岗位数据是否存在，若不存在则自动添加12个默认岗位"""
    if db.session.query(Job).count() > 0:
        return
    
    # 12个默认岗位数据
    jobs_data = [
        {
            "title": "Java后端开发工程师",
            "industry": "计算机软件、互联网、IT服务",
            "salary": "15000-25000",
            "work_address": "北京、上海、深圳、杭州、成都",
            "job_desc": "掌握 Java 基础、集合、并发、JVM、Spring Boot、Spring MVC、MyBatis/MySQL、Redis、Linux、Git。能把业务需求抽象成模块，能做接口重构、性能优化、缓存设计。",
            "requirements_text": "掌握 Java 语法、集合、多线程、JVM 基础；Spring/Spring Boot/MyBatis；数据库（MySQL）、缓存（Redis）、消息队列；LeetCode 基础题",
            "job_profile": {
                "skills": ["Java", "Spring Boot", "Spring MVC", "MyBatis", "MySQL", "Redis", "Linux", "Git"],
                "certificates": ["软考中级", "蓝桥杯", "英语四级"],
                "abilities": {
                    "创新能力": 4,
                    "学习能力": 5,
                    "抗压能力": 4,
                    "沟通能力": 4,
                    "实习能力": 4
                },
                "experience": "有后台管理系统、权限系统、电商/ERP/教务类项目经验优先"
            }
        },
        {
            "title": "前端开发工程师",
            "industry": "计算机软件、互联网、IT服务",
            "salary": "12000-20000",
            "work_address": "北京、上海、深圳、杭州、成都",
            "job_desc": "掌握 HTML5、CSS3、JavaScript/TypeScript，熟悉 Vue 或 React，了解组件化、工程化、接口联调与性能优化。能封装公共组件，改进页面交互、首屏性能和用户体验。",
            "requirements_text": "HTML/CSS/JavaScript 基础；至少一种主流框架（React/Vue）；Webpack/Vite 等工程化；性能优化与常见浏览器兼容问题",
            "job_profile": {
                "skills": ["HTML5", "CSS3", "JavaScript", "TypeScript", "Vue", "React", "Webpack", "Vite"],
                "certificates": ["作品集", "GitHub 项目", "蓝桥杯"],
                "abilities": {
                    "创新能力": 4,
                    "学习能力": 5,
                    "抗压能力": 3,
                    "沟通能力": 4,
                    "实习能力": 4
                },
                "experience": "有中后台系统、可视化大屏、小程序或官网项目经验优先"
            }
        },
        {
            "title": "C/C++开发工程师",
            "industry": "计算机软件、电子/半导体/集成电路、通信/网络设备",
            "salary": "15000-25000",
            "work_address": "北京、上海、深圳、成都、西安",
            "job_desc": "掌握 C/C++ 语法、数据结构、指针/内存管理、多线程、Linux 开发，了解网络编程、Qt、底层调试或嵌入式更佳。能围绕性能、稳定性、资源占用进行持续优化。",
            "requirements_text": "C/C++ 语法、数据结构、指针/内存管理、多线程、Linux 开发，了解网络编程、Qt、底层调试或嵌入式",
            "job_profile": {
                "skills": ["C++", "数据结构", "指针/内存管理", "多线程", "Linux 开发", "网络编程", "Qt", "嵌入式"],
                "certificates": ["CCPC/ICPC", "软考", "英语四级"],
                "abilities": {
                    "创新能力": 4,
                    "学习能力": 4,
                    "抗压能力": 5,
                    "沟通能力": 3,
                    "实习能力": 4
                },
                "experience": "有客户端、工业软件、驱动模块、底层组件或嵌入式项目经验优先"
            }
        },
        {
            "title": "软件测试工程师",
            "industry": "计算机软件、互联网、IT服务",
            "salary": "8000-15000",
            "work_address": "北京、上海、深圳、成都、西安",
            "job_desc": "掌握测试流程、测试用例设计、缺陷管理，了解 SQL、Linux、抓包、接口测试。能从业务流程中主动发现边界条件、异常场景和风险点。",
            "requirements_text": "测试流程、测试用例设计、缺陷管理，了解 SQL、Linux、抓包、接口测试",
            "job_profile": {
                "skills": ["测试流程", "测试用例设计", "缺陷管理", "SQL", "Linux", "抓包", "接口测试"],
                "certificates": ["ISTQB", "软考或测试相关竞赛"],
                "abilities": {
                    "创新能力": 3,
                    "学习能力": 4,
                    "抗压能力": 4,
                    "沟通能力": 4,
                    "实习能力": 4
                },
                "experience": "有 Web/App 测试、接口测试或驻场测试经验优先"
            }
        },
        {
            "title": "自动化测试工程师",
            "industry": "计算机软件、互联网、IT服务",
            "salary": "10000-18000",
            "work_address": "北京、上海、深圳、成都、西安",
            "job_desc": "熟悉 Python 或 Java，了解 Selenium、Playwright、Appium、接口自动化、持续集成和基础性能测试。能把重复测试场景沉淀为可复用脚本和自动化框架。",
            "requirements_text": "Python 或 Java，Selenium、Playwright、Appium、接口自动化、持续集成和基础性能测试",
            "job_profile": {
                "skills": ["Python", "Java", "Selenium", "Playwright", "Appium", "接口自动化", "持续集成", "性能测试"],
                "certificates": ["ISTQB", "高级测试证书"],
                "abilities": {
                    "创新能力": 4,
                    "学习能力": 4,
                    "抗压能力": 4,
                    "沟通能力": 3,
                    "实习能力": 4
                },
                "experience": "有自动化测试框架、CI 接入或平台化测试经验优先"
            }
        },
        {
            "title": "硬件测试工程师",
            "industry": "电子/半导体/集成电路、通信/网络设备、计算机硬件",
            "salary": "9000-16000",
            "work_address": "深圳、上海、北京、成都、西安",
            "job_desc": "熟悉硬件测试流程，掌握示波器、万用表、频谱仪等仪器使用，了解板卡、接口、可靠性/环境测试。能提出测试改进建议并辅助优化硬件设计。",
            "requirements_text": "硬件测试流程，示波器、万用表、频谱仪等仪器使用，板卡、接口、可靠性/环境测试",
            "job_profile": {
                "skills": ["硬件测试流程", "示波器", "万用表", "频谱仪", "板卡测试", "接口测试", "可靠性测试", "环境测试"],
                "certificates": ["电子设计竞赛", "英语四级"],
                "abilities": {
                    "创新能力": 3,
                    "学习能力": 4,
                    "抗压能力": 4,
                    "沟通能力": 3,
                    "实习能力": 4
                },
                "experience": "有实验室测试、板卡测试、电子设计或通信项目经验优先"
            }
        },
        {
            "title": "实施工程师",
            "industry": "IT服务、计算机软件、互联网",
            "salary": "7000-12000",
            "work_address": "全国",
            "job_desc": "了解软件部署、系统配置、数据库基础、数据导入导出、客户培训和实施文档编写，熟悉 Oracle/MySQL 更佳。能结合客户现场问题提出实施优化方案。",
            "requirements_text": "软件部署、系统配置、数据库基础、数据导入导出、客户培训和实施文档编写，熟悉 Oracle/MySQL",
            "job_profile": {
                "skills": ["软件部署", "系统配置", "数据库基础", "数据导入导出", "客户培训", "文档编写", "Oracle", "MySQL"],
                "certificates": ["软考", "PMP 基础认知", "行业软件/GIS 证书"],
                "abilities": {
                    "创新能力": 3,
                    "学习能力": 4,
                    "抗压能力": 5,
                    "沟通能力": 5,
                    "实习能力": 4
                },
                "experience": "有驻场实施、售后支持、项目交付或文档整理经验优先"
            }
        },
        {
            "title": "技术支持工程师",
            "industry": "IT服务、计算机软件、互联网",
            "salary": "7000-12000",
            "work_address": "全国",
            "job_desc": "熟悉产品部署、故障排查、系统维护、网络与数据库基础，能进行客户问题定位、培训和技术文档输出。能沉淀常见问题知识库，优化支持流程和响应机制。",
            "requirements_text": "产品部署、故障排查、系统维护、网络与数据库基础，客户问题定位、培训和技术文档输出",
            "job_profile": {
                "skills": ["产品部署", "故障排查", "系统维护", "网络基础", "数据库基础", "客户培训", "技术文档"],
                "certificates": ["软考", "厂商认证", "英语四级"],
                "abilities": {
                    "创新能力": 3,
                    "学习能力": 4,
                    "抗压能力": 4,
                    "沟通能力": 5,
                    "实习能力": 4
                },
                "experience": "有售后支持、驻场维护、Helpdesk 或项目支持经验优先"
            }
        },
        {
            "title": "AI科研工程师（算法/大数据方向）",
            "industry": "计算机软件、互联网、学术/科研",
            "salary": "20000-35000",
            "work_address": "北京、上海、深圳、杭州、成都",
            "job_desc": "掌握 Python、机器学习/深度学习基础，熟悉 PyTorch/TensorFlow，了解数据处理、模型训练、论文复现与实验记录。能提出实验改进思路，完成算法验证并形成研究结论。",
            "requirements_text": "Python、数据结构与算法、机器学习/深度学习框架（PyTorch/TensorFlow）；线性代数、概率统计；常见任务（分类/回归/推荐/NLP/CV）",
            "job_profile": {
                "skills": ["Python", "机器学习", "深度学习", "PyTorch", "TensorFlow", "数据处理", "模型训练", "论文复现"],
                "certificates": ["Kaggle/天池", "数学建模", "英语六级"],
                "abilities": {
                    "创新能力": 5,
                    "学习能力": 5,
                    "抗压能力": 4,
                    "沟通能力": 3,
                    "实习能力": 4
                },
                "experience": "有论文复现、竞赛、科研课题、模型部署或数据处理项目经验优先"
            }
        },
        {
            "title": "信息化项目助理",
            "industry": "IT服务、计算机软件、企业服务",
            "salary": "6000-10000",
            "work_address": "北京、上海、深圳、成都、西安",
            "job_desc": "了解项目管理流程，能进行进度跟踪、需求整理、会议纪要、文档输出与跨部门协调，熟悉 Office、Visio、Project 等工具。能优化模板、台账与项目协同流程，提高项目执行效率。",
            "requirements_text": "项目管理流程，进度跟踪、需求整理、会议纪要、文档输出与跨部门协调，Office、Visio、Project 等工具",
            "job_profile": {
                "skills": ["项目管理", "进度跟踪", "需求整理", "会议纪要", "文档输出", "Office", "Visio", "Project"],
                "certificates": ["PMP 基础认知", "软考", "英语四级"],
                "abilities": {
                    "创新能力": 3,
                    "学习能力": 4,
                    "抗压能力": 4,
                    "沟通能力": 5,
                    "实习能力": 4
                },
                "experience": "有 PMO、项目助理、实施支持、招投标文档或需求整理经验优先"
            }
        },
        {
            "title": "Java开发实习生",
            "industry": "计算机软件、互联网、IT服务",
            "salary": "3000-6000",
            "work_address": "北京、上海、深圳、杭州、成都",
            "job_desc": "掌握 Java 基础、面向对象、数据库和 SQL，了解 Spring Boot、Git、Linux 常用命令。能主动优化小模块和重复工作。",
            "requirements_text": "Java 基础、面向对象、数据库和 SQL，了解 Spring Boot、Git、Linux 常用命令",
            "job_profile": {
                "skills": ["Java 基础", "面向对象", "数据库", "SQL", "Spring Boot", "Git", "Linux"],
                "certificates": ["蓝桥杯", "计算机二级", "英语四级"],
                "abilities": {
                    "创新能力": 2,
                    "学习能力": 5,
                    "抗压能力": 3,
                    "沟通能力": 3,
                    "实习能力": 5
                },
                "experience": "每周可稳定到岗，有课程项目、后台系统或接口开发经历优先"
            }
        },
        {
            "title": "实施工程师（实习）",
            "industry": "IT服务、计算机软件、互联网",
            "salary": "2500-5000",
            "work_address": "全国",
            "job_desc": "了解系统安装部署、Excel 数据处理、数据库基础、实施文档和客户培训流程。能总结常见现场问题并沉淀模板。",
            "requirements_text": "系统安装部署、Excel 数据处理、数据库基础、实施文档和客户培训流程",
            "job_profile": {
                "skills": ["系统安装部署", "Excel 数据处理", "数据库基础", "实施文档", "客户培训"],
                "certificates": ["英语四级", "软考基础"],
                "abilities": {
                    "创新能力": 2,
                    "学习能力": 4,
                    "抗压能力": 4,
                    "沟通能力": 4,
                    "实习能力": 5
                },
                "experience": "实习期稳定 3-6 个月，有校内项目、客服支持或实施辅助经历优先"
            }
        }
    ]
    
    # 添加岗位数据到数据库
    for job_data in jobs_data:
        job = Job(
            title=job_data['title'],
            industry=job_data['industry'],
            salary=job_data['salary'],
            requirements_text=job_data['requirements_text'],
            job_profile_json=json.dumps(job_data['job_profile'], ensure_ascii=False),
            work_address=job_data['work_address'],
            job_desc=job_data['job_desc']
        )
        db.session.add(job)
    
    db.session.commit()


def _get_ai() -> AIHelper:
    # 使用单例模式的AI客户端
    return AIHelper(client=AIClient())


@api.errorhandler(PermissionError)
def _handle_perm(e):
    return _json_error(str(e), status=401)


@api.route("/register", methods=["POST"])
def register():
    """用户注册"""
    payload = request.get_json(silent=True) or {}
    username = (payload.get("username") or "").strip()
    password = (payload.get("password") or "").strip()
    name = (payload.get("name") or "").strip()

    if not username or not password:
        return _json_error("用户名和密码不能为空", status=400)

    if len(username) < 3 or len(username) > 20:
        return _json_error("用户名长度应在3-20个字符之间", status=400)

    if len(password) < 6:
        return _json_error("密码长度至少为6个字符", status=400)

    # 检查用户名是否已存在
    existing_student = db.session.query(Student).filter_by(username=username).one_or_none()
    if existing_student:
        return _json_error("用户名已存在", status=400)

    # 创建新用户
    student = Student(
        username=username,
        password=_hash_password(password),
        name=name or username,
        major="计算机相关"
    )
    db.session.add(student)
    db.session.commit()

    return jsonify({"ok": True, "message": "注册成功，请登录"})


@api.route("/login", methods=["POST"])
def login():
    payload = request.get_json(silent=True) or {}
    username = (payload.get("username") or "").strip()
    password = (payload.get("password") or "").strip()

    # 查找用户
    student = db.session.query(Student).filter_by(username=username).one_or_none()
    if not student:
        # 对于admin用户，如果不存在则创建
        if username == "admin" and password == "123456":
            student = Student(
                username=username,
                password=_hash_password(password),
                name="管理员",
                major="计算机相关"
            )
            db.session.add(student)
            db.session.commit()
        else:
            return _json_error("用户名或密码错误", status=401)
    else:
        # 验证密码
        if student.password != _hash_password(password):
            return _json_error("用户名或密码错误", status=401)

    token = secrets.token_urlsafe(24)
    TOKENS[token] = student.id

    return jsonify({"ok": True, "token": token, "student_id": student.id, "student": student.to_dict()})


@api.route("/upload_resume", methods=["POST"])
def upload_resume():
    student = _require_student()
    ai = _get_ai()
    _seed_jobs_if_empty(ai)

    manual = {
        "major": (request.form.get("major") or "").strip(),
        "skills": (request.form.get("skills") or "").strip(),
        "certs": (request.form.get("certs") or "").strip(),
        "internships": (request.form.get("internships") or "").strip(),
    }
    resume_text = (request.form.get("resume_text") or "").strip()

    file = request.files.get("resume_file")
    if file and file.filename:
        filename = file.filename.lower()
        data = file.read()
        if filename.endswith(".docx"):
            resume_text = (resume_text + "\n" + _read_docx_bytes(data)).strip()
        elif filename.endswith(".pdf"):
            resume_text = (resume_text + "\n" + _read_pdf_bytes(data)).strip()
        elif filename.endswith(".jpg") or filename.endswith(".jpeg") or filename.endswith(".png"):
                # 处理图片简历
                from pathlib import Path
                temp_path = str(Path(__file__).parent / "temp_resume.jpg")
                with open(temp_path, "wb") as f:
                    f.write(data)
                
                try:
                    # 初始化OCR引擎（使用兼容版本的参数）
                    # ocr = PaddleOCR(lang='ch', use_angle_cls=True)
                    # 识别图片文本
                    # result = ocr.ocr(temp_path)
                    
                    # 提取文本
                    ocr_text = ""  # 暂时返回空字符串，因为OCR模块未安装
                    # if result:
                    #     # 适配PaddleOCR 2.6.1.3的返回格式
                    #     if isinstance(result, list):
                    #         for item in result:
                    #             if isinstance(item, list):
                    #                 for line in item:
                    #                     if isinstance(line, list) and len(line) > 1:
                    #                         text = line[1][0] if len(line[1]) > 0 else ""
                    #                         ocr_text += text + "\n"
                    
                    # 更新简历文本
                    if ocr_text:
                        resume_text = (resume_text + "\n" + ocr_text).strip()
                    else:
                        resume_text = (resume_text + "\n" + "（已上传图片简历：OCR解析结果为空，请在手动录入区补充关键信息。）").strip()
                except Exception as e:
                    # OCR失败时的兜底处理
                    error_msg = str(e)[:100]  # 限制错误信息长度
                    resume_text = (resume_text + "\n" + f"（已上传图片简历：OCR解析失败 [{error_msg}]，请在手动录入区补充关键信息。）").strip()
                finally:
                    # 清理临时文件
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
        else:
            return _json_error("仅支持 docx/pdf/jpg/png 文件")

    if not resume_text:
        return _json_error("请上传简历文件或手动录入简历信息")

    parsed = ai.parse_resume(resume_text, manual=manual)

    if manual.get("major"):
        student.major = manual["major"]
    student.resume_raw_text = resume_text
    student.resume_parsed_json = _safe_json_dumps(parsed)
    student.increment_version()  # 增加数据版本号
    db.session.commit()

    # 缓存会自动失效，因为版本号已更新
    # 不需要手动清除缓存，新的缓存键会自动生成
    
    # 轻量级记忆：画像摘要，供 Agent 上下文联系
    summary = f"专业:{student.major or '未知'}；技能:{','.join((parsed.get('skills') or [])[:8])}；竞争力:{parsed.get('competitiveness_score','—')}"
    set_profile_summary(student.id, summary)

    # 写入一条 AI 引导消息（回到 chat.html 不丢失）
    db.session.add(
        ChatMessage(
            student_id=student.id,
            role="ai",
            content="简历解析完成！你的个人能力画像已生成，点击【性格测评】完善你的职业偏好吧～",
        )
    )
    db.session.commit()

    return jsonify({"ok": True, "student_id": student.id, "resume_parsed": parsed})


@api.route("/submit_test", methods=["POST"])
def submit_test():
    student = _require_student()
    payload = request.get_json(silent=True) or {}
    student.personality_test_json = _safe_json_dumps(payload)
    student.increment_version()  # 增加数据版本号
    db.session.commit()

    # 性格分析 → 短板诊断 + 岗位适配，写入轻量级记忆（上下文联系）
    try:
        resume_profile = json.loads(student.resume_parsed_json or "{}")
    except Exception:
        resume_profile = {}
    personality_test = payload
    ai = _get_ai()
    diag = ai.analyze_personality_for_jobs(resume_profile, personality_test)
    set_personality_diagnosis(student.id, diag)

    db.session.add(
        ChatMessage(
            student_id=student.id,
            role="ai",
            content="职业偏好已完善，我会结合你的简历和测评做精准人岗匹配，结果会在【岗位列表】中标注匹配度哦～",
        )
    )
    db.session.commit()

    # 清除相关缓存
    from cache import cache, get_cache_key
    # 清除匹配结果列表缓存
    cache_key = get_cache_key("match_jobs_results", student.id)
    cache.delete(cache_key)
    # 清除所有与该学生相关的匹配结果缓存
    # 注意：这里简化处理，实际应用中可能需要更精细的缓存管理
    
    return jsonify({"ok": True, "student_id": student.id, "personality_diagnosis": diag})


@api.route("/jobs", methods=["GET"])
def get_jobs():
    student: Optional[Student] = None
    try:
        student = _require_student()
    except Exception:
        student = None

    ai = _get_ai()
    _seed_jobs_if_empty(ai)

    q = (request.args.get("q") or "").strip().lower()
    include_match = (request.args.get("include_match") or "").strip() in {"1", "true", "yes"}
    page = int(request.args.get("page", 1))
    page_size = int(request.args.get("page_size", 20))
    page_size = min(page_size, 50)  # 限制每页最大数量

    # 构建查询
    query = db.session.query(Job)
    if q:
        like = f"%{q}%"
        query = query.filter((Job.title.like(like)) | (Job.industry.like(like)) | (Job.requirements_text.like(like)))

    # 计算总数
    total = query.count()
    
    # 分页查询
    offset = (page - 1) * page_size
    jobs = query.order_by(Job.id.asc()).offset(offset).limit(page_size).all()
    out = []

    student_profile = None
    if include_match and student and student.resume_parsed_json:
        try:
            student_profile = json.loads(student.resume_parsed_json)
            # 添加student_id到profile中，以便缓存
            student_profile['id'] = student.id
        except Exception:
            student_profile = None

    # 批量处理岗位
    jobs_to_update = []
    for j in jobs:
        item = j.to_dict()
        item['id'] = j.id  # 添加id字段
        
        if include_match and student_profile:
            # 尝试从缓存获取匹配结果
            from cache import get_cached_match_result
            cached_match = get_cached_match_result(student.id, j.id)
            if cached_match:
                item["match_preview"] = {"score": cached_match.get("score", 0), "reasoning": cached_match.get("reasoning", "")}
            else:
                try:
                    job_profile = json.loads(j.job_profile_json or "{}")
                    job_profile['id'] = j.id  # 添加id字段
                except Exception:
                    job_profile = ai.generate_job_profile(item)
                    j.job_profile_json = _safe_json_dumps(job_profile)
                    jobs_to_update.append(j)
                
                m = ai.match(student_profile, job_profile)
                item["match_preview"] = {"score": m.get("score", 0), "reasoning": m.get("reasoning", "")}
        
        out.append(item)

    # 批量更新岗位画像
    if jobs_to_update:
        db.session.bulk_save_objects(jobs_to_update)
        db.session.commit()

    return jsonify({"ok": True, "jobs": out, "total": total, "page": page, "page_size": page_size})


@api.route("/jobs/<int:job_id>", methods=["GET"])
def job_detail(job_id: int):
    job = db.session.get(Job, job_id)
    if not job:
        return _json_error("岗位不存在", status=404)
    
    # 尝试获取学生信息，计算匹配度
    try:
        student = _require_student()
        if student.resume_parsed_json:
            try:
                student_profile = json.loads(student.resume_parsed_json)
                student_profile['id'] = student.id
                
                # 尝试从缓存获取匹配结果
                cached_match = get_cached_match_result(student.id, job_id)
                if cached_match:
                    matched = cached_match
                else:
                    # 尝试从缓存获取岗位画像
                    cached_job_profile = get_cached_job_profile(job_id)
                    if cached_job_profile:
                        job_profile = cached_job_profile
                    else:
                        try:
                            job_profile = json.loads(job.job_profile_json or "{}")
                            job_profile['id'] = job_id
                        except Exception:
                            ai = _get_ai()
                            job_dict = job.to_dict()
                            job_dict['id'] = job_id
                            job_profile = ai.generate_job_profile(job_dict)
                            job.job_profile_json = _safe_json_dumps(job_profile)
                            db.session.add(job)
                            db.session.commit()
                        # 缓存岗位画像
                        cache_job_profile(job_id, job_profile)
                    
                    # 计算匹配结果
                    ai = _get_ai()
                    matched = ai.match(student_profile, job_profile)
                    # 缓存匹配结果
                    cache_match_result(student.id, job_id, matched)
                
                job_data = job.to_dict()
                job_data['match_result'] = matched
                return jsonify({"ok": True, "job": job_data})
            except Exception:
                pass
    except Exception:
        pass
    
    return jsonify({"ok": True, "job": job.to_dict()})


@api.route("/me", methods=["GET"])
def me():
    student = _require_student()
    data = student.to_dict()
    try:
        data["resume_parsed"] = json.loads(student.resume_parsed_json or "{}")
    except Exception:
        data["resume_parsed"] = {}
    try:
        data["personality_test"] = json.loads(student.personality_test_json or "{}")
    except Exception:
        data["personality_test"] = {}
    return jsonify({"ok": True, "student": data})


@api.route("/match_jobs", methods=["POST"])
def match_jobs():
    student = _require_student()
    ai = _get_ai()
    _seed_jobs_if_empty(ai)

    payload = request.get_json(silent=True) or {}
    goal = (payload.get("goal") or "").strip()

    if not student.resume_parsed_json:
        return _json_error("请先上传/录入简历并解析")

    try:
        student_profile = json.loads(student.resume_parsed_json)
        # 添加student_id到profile中，以便缓存
        student_profile['id'] = student.id
    except Exception:
        return _json_error("简历解析数据损坏，请重新上传解析")

    # 尝试从缓存获取匹配结果列表
    from cache import get_cache_key, cache
    # 使用学生数据版本号作为缓存键的一部分，确保数据更新时缓存自动失效
    cache_key = get_cache_key("match_jobs_results", student.id, student.data_version)
    cached_results = cache.get(cache_key)
    if cached_results:
        return jsonify({"ok": True, "student_id": student.id, "goal": goal, "matches": cached_results, "from_cache": True})

    # 批量获取所有岗位
    jobs = db.session.query(Job).order_by(Job.id.asc()).all()
    results = []
    match_results = []
    jobs_to_update = []

    # 清理旧匹配（演示版：每次重算覆盖）
    db.session.query(MatchResult).filter(MatchResult.student_id == student.id).delete()
    # 不立即提交，等待批量操作

    for j in jobs:
        # 尝试从缓存获取匹配结果
        cached_match = get_cached_match_result(student.id, j.id)
        if cached_match:
            matched = cached_match
        else:
            job_dict = j.to_dict()
            # 添加job_id到dict中，以便缓存
            job_dict['id'] = j.id
            
            # 尝试从缓存获取岗位画像
            cached_job_profile = get_cached_job_profile(j.id)
            if cached_job_profile:
                job_profile = cached_job_profile
            else:
                try:
                    job_profile = json.loads(j.job_profile_json or "{}")
                    # 添加job_id到profile中，以便缓存
                    job_profile['id'] = j.id
                except Exception:
                    job_profile = ai.generate_job_profile(job_dict)
                    j.job_profile_json = _safe_json_dumps(job_profile)
                    jobs_to_update.append(j)
                # 缓存岗位画像
                cache_job_profile(j.id, job_profile)

            # 计算匹配结果
            matched = ai.match(student_profile, job_profile)
            # 缓存匹配结果
            cache_match_result(student.id, j.id, matched)

        dim_full = matched.get("dimension_scores") or {}
        dim_4 = matched.get("dimension_scores_4") or {}
        mr = MatchResult(
            student_id=student.id,
            job_id=j.id,
            score=float(matched.get("score") or 0),
            dimension_scores_json=_safe_json_dumps({"dimension_scores": dim_full, "dimension_scores_4": dim_4}),
            gap_analysis_json=_safe_json_dumps(matched.get("gap_analysis") or {}),
            reasoning=matched.get("reasoning") or "",
        )
        match_results.append(mr)
        results.append(
            {
                "job_id": j.id,
                "job_title": j.title,
                "industry": j.industry,
                "salary": j.salary,
                "score": mr.score,
                "dimension_scores": dim_full,
                "dimension_scores_4": dim_4,
                "gap_analysis": matched.get("gap_analysis") or {},
                "reasoning": mr.reasoning,
                "online_result": matched.get("online_result"),  # 在线结果
                "offline_result": matched.get("offline_result")  # 离线结果
            }
        )

    # 批量添加和更新
    if jobs_to_update:
        db.session.bulk_save_objects(jobs_to_update)
    if match_results:
        db.session.bulk_save_objects(match_results)
    db.session.commit()

    # 排序结果
    results.sort(key=lambda x: x["score"], reverse=True)
    
    # 缓存匹配结果列表，使用版本号作为缓存键的一部分
    cache_key = get_cache_key("match_jobs_results", student.id, student.data_version)
    cache.set(cache_key, results)

    return jsonify({"ok": True, "student_id": student.id, "goal": goal, "matches": results, "from_cache": False})


@api.route("/generate_report", methods=["POST"])
def generate_report():
    student = _require_student()
    ai = _get_ai()
    _seed_jobs_if_empty(ai)

    payload = request.get_json(silent=True) or {}
    goal = (payload.get("goal") or "").strip()

    if not student.resume_parsed_json:
        return _json_error("请先上传/录入简历并解析")

    # 取已有匹配或先计算
    match_rows = (
        db.session.query(MatchResult)
        .filter(MatchResult.student_id == student.id)
        .order_by(MatchResult.score.desc())
        .limit(5)
        .all()
    )
    if not match_rows:
        # 自动计算一次
        _ = match_jobs()
        match_rows = (
            db.session.query(MatchResult)
            .filter(MatchResult.student_id == student.id)
            .order_by(MatchResult.score.desc())
            .limit(5)
            .all()
        )

    top_matches = []
    for mr in match_rows:
        job = db.session.get(Job, mr.job_id)
        try:
            dims = json.loads(mr.dimension_scores_json or "{}")
            dim_scores = dims.get("dimension_scores") or dims
            dim_scores_4 = dims.get("dimension_scores_4") or {}
        except Exception:
            dim_scores = {}
            dim_scores_4 = {}
        top_matches.append(
            {
                "job_id": mr.job_id,
                "job_title": job.title if job else f"job#{mr.job_id}",
                "score": mr.score,
                "gap_analysis": json.loads(mr.gap_analysis_json or "{}"),
                "dimension_scores": dim_scores,
                "dimension_scores_4": dim_scores_4,
                "reasoning": mr.reasoning,
            }
        )

    student_dict = student.to_dict()
    content = ai.generate_report(student_dict, top_matches, goal=goal)

    from ai_helper import report_completeness_check
    check = report_completeness_check(content)

    report = Report(student_id=student.id, content=content)
    db.session.add(report)
    db.session.commit()

    return jsonify(
        {
            "ok": True,
            "report_id": report.id,
            "content": content,
            "download_url": f"/download_report/{report.id}",
            "download_docx_url": f"/download_report_docx/{report.id}",
            "download_pdf_url": f"/download_report_pdf/{report.id}",
            "completeness_check": check,
            "created_at": datetime.utcnow().isoformat(),
        }
    )


@api.route("/career_graph", methods=["GET"])
def career_graph():
    """岗位关联图谱：垂直晋升 + 换岗路径。至少 5 个岗位族，每族不少于 2 条换岗路径。"""
    title = (request.args.get("title") or "").strip()
    dynamic = (request.args.get("dynamic") or "").strip() in {"1", "true", "yes"}
    
    if dynamic:
        # 尝试获取学生技能和兴趣
        student_skills = []
        interests = []
        try:
            student = _require_student()
            if student.resume_parsed_json:
                import json
                resume_data = json.loads(student.resume_parsed_json)
                student_skills = resume_data.get("skills", [])
            if student.personality_test_json:
                import json
                test_data = json.loads(student.personality_test_json)
                interests = test_data.get("career_interest", {}).get("interests", [])
        except Exception:
            pass
        
        data = build_dynamic_echarts_data(title, student_skills, interests)
    else:
        data = build_echarts_data(title)
    
    return jsonify({"ok": True, "graph": data, "families": list_all_families()})


@api.route("/feedback", methods=["POST"])
def submit_feedback():
    """提交用户反馈"""
    student = _require_student()
    payload = request.get_json(silent=True) or {}
    
    job_id = int(payload.get("job_id") or 0)
    rating = int(payload.get("rating") or 0)
    comment = (payload.get("comment") or "").strip()
    feedback_type = (payload.get("feedback_type") or "match").strip()
    
    if not job_id:
        return _json_error("缺少 job_id")
    
    if rating < 1 or rating > 5:
        return _json_error("评分必须在1-5之间")
    
    job = db.session.get(Job, job_id)
    if not job:
        return _json_error("岗位不存在", status=404)
    
    # 创建反馈记录
    feedback = Feedback(
        student_id=student.id,
        job_id=job_id,
        rating=rating,
        comment=comment,
        feedback_type=feedback_type
    )
    db.session.add(feedback)
    db.session.commit()
    
    # 这里可以添加基于反馈的匹配算法优化逻辑
    # 例如，调整技能权重或维度权重
    
    return jsonify({"ok": True, "feedback_id": feedback.id})


@api.route("/feedback/history", methods=["GET"])
def get_feedback_history():
    """获取用户反馈历史"""
    student = _require_student()
    
    feedbacks = db.session.query(Feedback).filter(
        Feedback.student_id == student.id
    ).order_by(Feedback.created_at.desc()).limit(50).all()
    
    result = []
    for feedback in feedbacks:
        job = db.session.get(Job, feedback.job_id)
        result.append({
            "id": feedback.id,
            "job_id": feedback.job_id,
            "job_title": job.title if job else "未知岗位",
            "rating": feedback.rating,
            "comment": feedback.comment,
            "feedback_type": feedback.feedback_type,
            "created_at": feedback.created_at.isoformat()
        })
    
    return jsonify({"ok": True, "feedbacks": result})


@api.route("/chat/history", methods=["GET"])
def chat_history():
    student = _require_student()
    limit = int(request.args.get("limit") or 50)
    rows = (
        db.session.query(ChatMessage)
        .filter(ChatMessage.student_id == student.id)
        .order_by(ChatMessage.id.asc())
        .limit(max(1, min(limit, 300)))
        .all()
    )
    return jsonify({"ok": True, "messages": [r.to_dict() for r in rows]})


def _offline_agent_reply(student: Student, user_text: str) -> str:
    # 离线兜底：给出引导+可执行建议（不联网）
    if not student.resume_parsed_json:
        return "我可以先帮你解析简历生成能力画像。请点击顶部【简历上传】上传/录入简历信息。"
    if not student.personality_test_json:
        return "我已拿到你的简历画像。建议再做一次【性格测评】补充职业偏好，我会基于两者做更精准的人岗匹配。"
    t = (user_text or "").strip()
    if "报告" in t:
        return "你可以点击顶部【报告生成】一键生成职业生涯发展报告（含目标与行动计划），并支持下载。"
    if "岗位" in t or "匹配" in t:
        return "我建议你先在【岗位列表】查看匹配度较高的岗位，并进入【匹配结果】页查看 TOP5、雷达图与差距清单。"
    return "你可以问我：市场岗位能力要求、我的优势与短板、推荐岗位与差距、以及分阶段行动计划。"


@api.route("/chat/send", methods=["POST"])
def chat_send():
    student = _require_student()
    payload = request.get_json(silent=True) or {}
    user_text = (payload.get("text") or "").strip()
    if not user_text:
        return _json_error("消息不能为空")

    # 存用户消息
    db.session.add(ChatMessage(student_id=student.id, role="user", content=user_text))
    db.session.commit()

    ai = _get_ai()
    client = ai.client

    # 轻量级记忆：先用 LLM 做一句话关键信息摘要，失败则用关键词匹配兜底
    try:
        summary = client.chat_text(
            "你是信息抽取助手，请用不超过一句话的中文提炼这段话的关键信息。",
            f"用户原话：{user_text}",
        )
        if summary:
            add_key_fact(student.id, summary.strip())
    except Exception:
        key_phrases = ["成绩", "内向", "外向", "压力", "家庭", "考研", "保研", "大厂", "国企", "实习", "项目", "证书", "迷茫", "焦虑"]
        for p in key_phrases:
            if p in user_text:
                add_key_fact(student.id, f"用户提及：{user_text[:80]}...")
                break

    # 博弈模型：Regret Matching 计算最优路径（大厂/国企/保研）
    combined = {"dimensions": {}, "personality": {}}
    try:
        rp = json.loads(student.resume_parsed_json or "{}")
        combined["dimensions"] = rp.get("dimensions") or {}
        combined["competitiveness_score"] = rp.get("competitiveness_score") or 50
    except Exception:
        pass
    try:
        pt = json.loads(student.personality_test_json or "{}")
        s = pt.get("scores") or {}
        combined["personality"] = {k: v for k, v in s.items()}
        combined["career_interest"] = pt.get("career_interest") or {}
        combined["resources"] = pt.get("resources") or []
        combined["constraints"] = pt.get("constraints") or {}
    except Exception:
        pass

    # 即时偏好：更想稳定 / 不想出国 等（影响博弈参数）
    prefs: dict[str, bool] = {}
    if "稳定" in user_text or "体制" in user_text or "编制" in user_text:
        prefs["prefer_stable"] = True
    if "不想出国" in user_text or "不想海外" in user_text or "留在国内" in user_text:
        prefs["avoid_abroad"] = True
    combined["user_prefs"] = prefs
    try:
        regret_res = regret_matching_from_profile(combined)
        set_regret_result(student.id, regret_res)
    except Exception:
        regret_res = None

    # 拼接上下文：记忆模块 + 博弈结果 + 学生画像 + 最近对话
    history_rows = (
        db.session.query(ChatMessage)
        .filter(ChatMessage.student_id == student.id)
        .order_by(ChatMessage.id.desc())
        .limit(20)
        .all()
    )
    history = list(reversed(history_rows))

    # 简易联网搜索 + 本地知识库：当用户询问“市场要求/竞赛/趋势/怎么学”等，补充检索摘要
    search_block = ""
    keywords = ["市场", "趋势", "要求", "竞赛", "证书", "学习路线", "怎么学", "什么是", "是什么"]
    if any(k in user_text for k in keywords):
        results = client.web_search(user_text, limit=5) if hasattr(client, "web_search") else []
        if results:
            lines = [f"- {r.get('title','')}（{r.get('url','')}）" for r in results]
            search_block = "【联网检索参考】\n" + "\n".join(lines) + "\n"
        # 本地知识库（轻量 RAG）
        local_snips = search_knowledge(user_text, max_chunks=3)
        if local_snips:
            ls = [s.get("snippet", "") for s in local_snips]
            search_block += "【本地知识库参考】\n" + "\n---\n".join(ls[:3]) + "\n"

    system_prompt = (
        "你是温暖、专业的大学生 AI 职业规划助手，像一位关心你的学长/学姐。"
        "你的目标：1) 帮学生快速了解就业市场/岗位能力要求；2) 准确分析就业能力与意愿；"
        "3) 给出明确可操作的建议。要求：语气亲切、有共情，同时专业、可执行；"
        "若系统提供了【性格与短板诊断】或【博弈路径建议】，务必结合这些信息综合给出最优建议。"
        "输出中文。\n"
    )

    memory_block = to_context_string(student.id)
    regret_block = ""
    if regret_res:
        regret_block = f"\n【博弈路径建议（Regret Matching）】最优路径：{regret_res.best_action}；{regret_res.recommendation}"

    context = {
        "student": student.to_dict(),
        "resume_parsed_json": student.resume_parsed_json,
        "personality_test_json": student.personality_test_json,
        "search": search_block,
        "recent_messages": [{"role": m.role, "content": m.content} for m in history],
    }
    user_prompt = (
        f"{search_block}"
        f"{memory_block}\n{regret_block}\n\n"
        f"【学生信息与画像】\n{_safe_json_dumps(context)}\n\n"
        f"【本轮用户问题】\n{user_text}\n\n"
        "请结合记忆、性格诊断、博弈路径建议与对话历史，亲切、专业地回答，并给出下一步操作引导。"
    )

    reply = ""
    try:
        reply = client.chat_text(system_prompt, user_prompt)
    except Exception:
        reply = _offline_agent_reply(student, user_text)

    db.session.add(ChatMessage(student_id=student.id, role="ai", content=reply))
    db.session.commit()

    return jsonify({"ok": True, "reply": reply})


@api.route("/download_report/<int:report_id>", methods=["GET"])
def download_report(report_id: int):
    student = _require_student()
    report = db.session.get(Report, report_id)
    if not report or report.student_id != student.id:
        return _json_error("报告不存在", status=404)

    data = report.content.encode("utf-8")
    return send_file(
        io.BytesIO(data),
        as_attachment=True,
        download_name=f"career_report_{report_id}.txt",
        mimetype="text/plain; charset=utf-8",
    )


@api.route("/download_report_docx/<int:report_id>", methods=["GET"])
def download_report_docx(report_id: int):
    student = _require_student()
    report = db.session.get(Report, report_id)
    if not report or report.student_id != student.id:
        return _json_error("报告不存在", status=404)

    doc = DocxDocument()
    for line in (report.content or "").splitlines():
        doc.add_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return send_file(
        bio,
        as_attachment=True,
        download_name=f"career_report_{report_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@api.route("/download_report_pdf/<int:report_id>", methods=["GET"])
def download_report_pdf(report_id: int):
    student = _require_student()
    report = db.session.get(Report, report_id)
    if not report or report.student_id != student.id:
        return _json_error("报告不存在", status=404)

    # 使用reportlab生成PDF文件，支持中文
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.fonts import addMapping
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    # 尝试添加中文字体支持
    try:
        # 尝试使用系统默认的中文字体
        # 对于Windows系统，通常有SimHei字体
        pdfmetrics.registerFont(TTFont('SimHei', 'simhei.ttf'))
        addMapping('SimHei', 0, 0, 'SimHei')
        addMapping('SimHei', 0, 1, 'SimHei')
        addMapping('SimHei', 1, 0, 'SimHei')
        addMapping('SimHei', 1, 1, 'SimHei')
    except Exception:
        # 如果没有找到SimHei字体，使用reportlab默认字体
        pass
    
    # 创建字节流
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4)
    story = []
    
    # 获取默认样式
    styles = getSampleStyleSheet()
    normal_style = styles['Normal']
    normal_style.alignment = TA_LEFT
    
    # 尝试设置中文字体
    try:
        normal_style.fontName = 'SimHei'
    except Exception:
        # 如果没有中文字体，使用默认字体
        pass
    
    # 写入内容
    content = report.content or ""
    for line in content.splitlines():
        if line.strip():
            para = Paragraph(line, normal_style)
            story.append(para)
        story.append(Spacer(1, 12))
    
    # 构建PDF
    doc.build(story)
    bio.seek(0)
    
    return send_file(
        bio,
        as_attachment=True,
        download_name=f"career_report_{report_id}.pdf",
        mimetype="application/pdf",
    )


@api.route("/reports/<int:report_id>", methods=["PUT"])
def update_report(report_id: int):
    """在线编辑后保存报告内容（覆盖原 content）。"""
    student = _require_student()
    payload = request.get_json(silent=True) or {}
    content = (payload.get("content") or "").strip()
    if not content:
        return _json_error("报告内容不能为空")

    report = db.session.get(Report, report_id)
    if not report or report.student_id != student.id:
        return _json_error("报告不存在或无权限", status=404)

    report.content = content
    db.session.commit()
    return jsonify({"ok": True, "report_id": report.id})


@api.route("/admin/users", methods=["GET"])
def admin_users():
    """管理员查看所有用户"""
    _require_admin()
    
    students = db.session.query(Student).all()
    users = [student.to_dict_admin() for student in students]
    
    return jsonify({"ok": True, "users": users})


@api.route("/admin/stats", methods=["GET"])
def admin_stats():
    """管理员查看统计信息"""
    _require_admin()
    
    # 统计基础数据
    user_count = db.session.query(Student).count()
    job_count = db.session.query(Job).count()
    match_count = db.session.query(MatchResult).count()
    
    # 统计岗位热度
    job_stats = db.session.query(
        Job.id,
        Job.title,
        db.func.count(MatchResult.id).label('match_count')
    ).outerjoin(
        MatchResult, Job.id == MatchResult.job_id
    ).group_by(
        Job.id, Job.title
    ).order_by(
        db.func.count(MatchResult.id).desc()
    ).limit(10).all()
    
    hot_jobs = [
        {"job_id": job.id, "job_title": job.title, "match_count": job.match_count}
        for job in job_stats
    ]
    
    # 统计技能缺口
    # 这里需要从MatchResult的gap_analysis_json中提取技能缺口信息
    # 由于数据存储在JSON中，我们需要先获取所有匹配结果，然后解析
    match_results = db.session.query(MatchResult).all()
    skill_gaps = {}
    
    for mr in match_results:
        try:
            gap_analysis = json.loads(mr.gap_analysis_json or "{}")
            missing_skills = gap_analysis.get("missing_skills", [])
            for skill in missing_skills:
                if isinstance(skill, str):
                    skill_gaps[skill] = skill_gaps.get(skill, 0) + 1
        except Exception:
            pass
    
    # 排序技能缺口
    sorted_skill_gaps = sorted(skill_gaps.items(), key=lambda x: x[1], reverse=True)[:10]
    skill_gap_stats = [
        {"skill": skill, "gap_count": count}
        for skill, count in sorted_skill_gaps
    ]
    
    return jsonify({
        "ok": True,
        "user_count": user_count,
        "job_count": job_count,
        "match_count": match_count,
        "hot_jobs": hot_jobs,
        "skill_gaps": skill_gap_stats
    })


@api.route("/admin/user/<int:user_id>/matches", methods=["GET"])
def admin_user_matches(user_id: int):
    """管理员查看指定用户的岗位匹配信息"""
    _require_admin()
    
    # 检查用户是否存在
    student = db.session.get(Student, user_id)
    if not student:
        return _json_error("用户不存在", status=404)
    
    # 获取用户的匹配结果
    match_results = db.session.query(MatchResult).filter(
        MatchResult.student_id == user_id
    ).order_by(
        MatchResult.score.desc()
    ).all()
    
    matches = []
    for mr in match_results:
        job = db.session.get(Job, mr.job_id)
        if job:
            matches.append({
                "match_id": mr.id,
                "job_id": mr.job_id,
                "job_title": job.title,
                "score": mr.score,
                "reasoning": mr.reasoning,
                "created_at": mr.created_at.isoformat()
            })
    
    return jsonify({
        "ok": True,
        "user": student.to_dict_admin(),
        "matches": matches
    })


@api.route("/admin/jobs", methods=["POST"])
def admin_add_job():
    """管理员添加岗位"""
    _require_admin()
    
    data = request.get_json()
    title = data.get("title")
    salary_range = data.get("salary_range")
    company = data.get("company", "默认公司")
    description = data.get("description", "岗位描述")
    requirements = data.get("requirements", "岗位要求")
    
    if not title:
        return _json_error("岗位名称不能为空", status=400)
    
    # 创建新岗位
    job = Job(
        title=title,
        company=company,
        salary_range=salary_range,
        description=description,
        requirements=requirements
    )
    
    db.session.add(job)
    db.session.commit()
    
    return jsonify({
        "ok": True,
        "message": "岗位添加成功",
        "job_id": job.id
    })


@api.route("/admin/jobs/<int:job_id>", methods=["DELETE"])
def admin_delete_job(job_id: int):
    """管理员删除岗位"""
    _require_admin()
    
    # 检查岗位是否存在
    job = db.session.get(Job, job_id)
    if not job:
        return _json_error("岗位不存在", status=404)
    
    # 删除相关的匹配结果
    db.session.query(MatchResult).filter(MatchResult.job_id == job_id).delete()
    
    # 删除岗位
    db.session.delete(job)
    db.session.commit()
    
    return jsonify({
        "ok": True,
        "message": "岗位删除成功"
    })


def _build_learning_plan_from_gap(
    student_profile: dict[str, Any],
    job: Job,
    job_profile: dict[str, Any],
    match_result: dict[str, Any],
    weekly_hours: int = 8,
    prefer_internship: bool = True,
) -> dict[str, Any]:
    """纯本地规则版：根据差距生成分阶段学习计划与作品集建议。"""
    missing_skills = (match_result.get("gap_analysis") or {}).get("missing_skills") or []
    missing_skills = [s for s in missing_skills if isinstance(s, str)]
    top_missing = missing_skills[:6]

    # 作品集项目建议（按岗位族/技能做粗分类）
    title = (job.title or "").strip()
    family = (job_profile.get("archetype") or {}).get("family") or (job_profile.get("industry") or "") or ""
    t = (title + " " + family).lower()
    projects: list[dict[str, Any]] = []
    if any(k in t for k in ["前端", "react", "vue", "frontend"]):
        projects = [
            {"name": "中后台管理系统（含权限/表格/图表）", "deliverables": ["可部署链接", "README", "截图", "关键组件封装说明"]},
            {"name": "性能优化小项目（首屏/懒加载/缓存）", "deliverables": ["优化前后对比数据", "复现步骤", "代码仓库"]},
        ]
    elif any(k in t for k in ["测试", "qa", "set"]):
        projects = [
            {"name": "接口自动化测试框架（pytest/requests + 报告）", "deliverables": ["用例覆盖说明", "CI 接入示例", "测试报告"]},
            {"name": "Web 自动化回归（Playwright/Selenium）", "deliverables": ["脚本稳定性说明", "失败截图/日志", "README"]},
        ]
    elif any(k in t for k in ["算法", "nlp", "llm", "ai", "科研", "大数据", "spark", "flink"]):
        projects = [
            {"name": "论文/模型复现 + 训练记录", "deliverables": ["复现实验日志", "指标对齐", "可复现脚本"]},
            {"name": "小型端到端应用（数据→训练→推理→服务）", "deliverables": ["API 服务", "Demo", "部署说明"]},
        ]
    elif any(k in t for k in ["实施", "技术支持", "交付", "项目"]):
        projects = [
            {"name": "部署交付演练（Docker + 文档 + 培训材料）", "deliverables": ["部署手册", "FAQ 知识库", "演示录屏/截图"]},
            {"name": "数据导入/报表/运维小工具", "deliverables": ["使用说明", "异常处理说明", "代码仓库"]},
        ]
    else:
        projects = [
            {"name": "通用作品集项目（与目标岗位JD对齐）", "deliverables": ["可部署链接", "README", "量化成果"]},
        ]

    # 分阶段计划（根据 weekly_hours 粗调强度）
    intensity = "轻量" if weekly_hours <= 6 else ("中等" if weekly_hours <= 12 else "高强度")
    short_weeks = 4 if weekly_hours <= 6 else 3
    mid_months = 3

    short_tasks = []
    if top_missing:
        short_tasks.append({"task": f"补齐关键技能 Top{min(3,len(top_missing))}", "items": top_missing[:3], "output": "学习笔记/小练习/可运行 demo"})
    short_tasks.append({"task": "准备 1 个可讲清楚的 STAR 案例", "items": ["项目背景", "个人职责", "技术方案", "量化结果"], "output": "面试讲稿 3 条"})
    if prefer_internship:
        short_tasks.append({"task": "实习投递闭环启动", "items": ["每周投递 30-50", "简历关键词对齐JD", "每周复盘一次"], "output": "投递台账 + 复盘记录"})

    mid_tasks = [
        {"task": "完成 1 个作品集项目并上线", "items": [p["name"] for p in projects[:1]], "output": "代码仓库 + README + 部署链接"},
        {"task": "补齐剩余缺口并做面试准备", "items": top_missing[3:6] if len(top_missing) > 3 else [], "output": "题单/总结/模拟面试记录"},
    ]

    plan = {
        "intensity": intensity,
        "weekly_hours": weekly_hours,
        "short_term": {
            "duration": f"{short_weeks}周",
            "goals": ["拉齐关键技能门槛", "形成可展示产出", "启动投递或准备节奏"],
            "tasks": short_tasks,
            "metrics": ["每周学习时长达标", "产出物≥2项", "模拟面试≥2次" if prefer_internship else "题单/复盘≥2次"],
        },
        "mid_term": {
            "duration": f"{mid_months}个月",
            "goals": ["作品集项目成型", "面试能力闭环", "匹配度明显提升"],
            "tasks": mid_tasks,
            "metrics": ["作品集项目≥1个可部署", "缺失技能清单减少≥50%", "面试邀约率提升（复盘可见）"],
        },
        "projects": projects,
    }
    return plan


@api.route("/plan_for_target_job", methods=["POST"])
def plan_for_target_job():
    """
    目标岗位提升计划：
    - 允许用户选择一个"目标岗位"（即使当前不匹配）
    - 返回差距分析 + 分阶段学习计划 + 作品集项目建议 +（可选）路径图谱摘要
    """
    student = _require_student()
    payload = request.get_json(silent=True) or {}
    job_id = int(payload.get("job_id") or 0)
    weekly_hours = int(payload.get("weekly_hours") or 8)
    weekly_hours = max(1, min(40, weekly_hours))
    prefer_internship = bool(payload.get("prefer_internship", True))

    if not job_id:
        return _json_error("缺少 job_id")

    job = db.session.get(Job, job_id)
    if not job:
        return _json_error("岗位不存在", status=404)

    if not student.resume_parsed_json:
        return _json_error("请先上传/录入简历并解析")

    try:
        student_profile = json.loads(student.resume_parsed_json)
        # 添加student_id到profile中，以便缓存
        student_profile['id'] = student.id
    except Exception:
        return _json_error("简历解析数据损坏，请重新上传解析")

    # 尝试从缓存获取学习计划
    cached_plan = get_cached_learning_plan(student.id, job_id)
    if cached_plan:
        # 检查缓存的计划是否包含所有必要信息
        if all(key in cached_plan for key in ["intensity", "weekly_hours", "short_term", "mid_term", "projects"]):
            # 从缓存获取匹配结果
            from cache import get_cached_match_result
            cached_match = get_cached_match_result(student.id, job_id)
            if cached_match:
                # 路径图谱摘要（用于"目标岗位→A→B→C"展示）
                graph = build_echarts_data(job.title)
                graph_summary = {
                    "family": graph.get("family"),
                    "nodes_count": len(graph.get("nodes") or []),
                    "links_count": len(graph.get("links") or []),
                }
                
                return jsonify(
                    {
                        "ok": True,
                        "job": job.to_dict(),
                        "match": cached_match,
                        "plan": cached_plan,
                        "career_graph_summary": graph_summary,
                        "from_cache": True
                    }
                )

    ai = _get_ai()
    job_dict = job.to_dict()
    job_dict['id'] = job.id  # 添加id字段
    
    try:
        job_profile = json.loads(job.job_profile_json or "{}")
        # 添加job_id到profile中，以便缓存
        job_profile['id'] = job.id
    except Exception:
        job_profile = ai.generate_job_profile(job_dict)
        # 更新岗位画像到数据库
        job.job_profile_json = _safe_json_dumps(job_profile)
        db.session.add(job)
        db.session.commit()

    # 复用匹配输出作为 gap 来源（即使不匹配，也能给缺口）
    # 提升计划生成时只使用离线匹配，提高性能
    match_result = ai.match(student_profile, job_profile, use_online=False)

    # 路径图谱摘要（用于"目标岗位→A→B→C"展示）
    graph = build_echarts_data(job.title)
    graph_summary = {
        "family": graph.get("family"),
        "nodes_count": len(graph.get("nodes") or []),
        "links_count": len(graph.get("links") or []),
    }

    plan = _build_learning_plan_from_gap(
        student_profile=student_profile,
        job=job,
        job_profile=job_profile,
        match_result=match_result,
        weekly_hours=weekly_hours,
        prefer_internship=prefer_internship,
    )

    # 缓存学习计划
    cache_learning_plan(student.id, job_id, plan)

    return jsonify(
        {
            "ok": True,
            "job": job_dict,
            "match": match_result,
            "plan": plan,
            "career_graph_summary": graph_summary,
            "from_cache": False
        }
    )